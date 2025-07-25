# -*- coding: utf-8 -*-
"""Untitled0.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1TpF5cIFgd4s8D2T3U_AfzOYFJWL52zE3
"""

from google.colab import drive
drive.mount('/content/drive')

# STEP 1: Mount Google Drive
from google.colab import drive
drive.mount('/content/drive')

# STEP 2: Install packages (add fuzzywuzzy for better string matching)
!pip install python-docx PyMuPDF wordcloud scikit-learn fuzzywuzzy python-Levenshtein
!pip install --upgrade scikit-learn

# STEP 3: Imports
import os
import re
import fitz # PyMuPDF
import docx
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from fuzzywuzzy import fuzz
from IPython.display import display, HTML # For rich output in Colab

# STEP 4: Path setup (ensure this folder exists in your Google Drive)
# Create a folder named 'resume_data' inside 'My Drive' and put your resumes there
folder_path = "/content/drive/MyDrive/resume data set/Resumes" # Updated folder path
if not os.path.exists(folder_path):
    # Changed message to reflect the user's provided path
    print(f"Error: Folder not found at {folder_path}. Please ensure this folder exists in your Google Drive and contains your resumes.")
    # Exit the script if the folder doesn't exist to prevent errors
    exit()


# STEP 5: Functions
def extract_text_from_pdf(path):
    text = ""
    try:
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF {path}: {e}")
    return text

def extract_text_from_docx(path):
    text = ""
    try:
        doc = docx.Document(path)
        text = " ".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX {path}: {e}")
    return text

def clean_text(text):
    # Remove non-ASCII characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    # Remove extra spaces and newlines
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    # Allow only alphanumeric, common punctuation, @ and .
    text = re.sub(r'[^a-zA-Z0-9@\s.,-]', '', text)
    return text.lower()

def extract_email(text):
    match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return match.group(0) if match else "N/A"

def extract_phone(text):
    # Broader pattern for phone numbers, including common separators and international codes
    # This pattern tries to capture 10-15 digit numbers, possibly with spaces, hyphens, or parentheses
    match = re.search(r'\+?\d{1,3}?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,9}\b', text)
    return match.group(0).replace(" ", "").replace("-", "").replace("(", "").replace(")", "") if match else "N/A"

def get_job_description():
    """Allows user to paste the job description directly with clear instructions."""
    print("\n" + "="*50)
    print("--- JOB DESCRIPTION INPUT ---")
    print("="*50)
    print("Please paste the full Job Description text below.")
    print("Type your job description line by line.")
    print("When you are finished, type 'END_JD' on a new, separate line and press Enter.")
    print("\n--- Start pasting your Job Description here: ---")

    lines = []
    while True:
        line = input() # This is where you type each line
        if line.strip().upper() == 'END_JD': # Case-insensitive check for 'END_JD'
            break
        lines.append(line)
    jd_text = "\n".join(lines)

    if not jd_text.strip(): # Check if JD is empty after stripping whitespace
        print("\nWarning: Job description was empty. Please try again.")
        return get_job_description() # Retry

    print("\n--- Job Description received. Processing... ---")
    return clean_text(jd_text)

def highlight_keywords_in_context(text, keywords, context_size=150):
    """Highlights keywords and provides a surrounding context."""
    highlighted_sections = []
    text_lower = text.lower() # Use lowercased text for matching

    for keyword in keywords:
        # Use a more robust pattern for whole words
        pattern = r'\b' + re.escape(keyword) + r'\b'
        for match in re.finditer(pattern, text_lower):
            start, end = match.span()
            # Find the start of the context
            context_start = max(0, start - context_size)
            # Find the end of the context
            context_end = min(len(text), end + context_size)

            # Ensure we're highlighting in the original case-preserved text
            original_match_text = text[start:end]

            # Extract the original case-preserved context
            context = text[context_start:context_end]

            # Highlight the specific matched keyword within this context for display
            # We use a non-capturing group (?:...) for the bold tag so it doesn't interfere
            highlighted_context = re.sub(r'\b' + re.escape(keyword) + r'\b', f'<b><span style="background-color: yellow;">{original_match_text}</span></b>', context, flags=re.IGNORECASE)

            highlighted_sections.append(f"...{highlighted_context}...")
    return "\n---\n".join(highlighted_sections) if highlighted_sections else "No significant skill matches found in context."


# STEP 6: Load and clean resumes
def load_and_clean_resumes(path):
    resume_files = [os.path.join(path, f) for f in os.listdir(path) if f.lower().endswith(('.pdf', '.docx'))]
    resume_texts = {}
    print(f"\n--- Loading {len(resume_files)} Resumes ---")
    for file in resume_files:
        print(f"Processing: {os.path.basename(file)}")
        if file.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file)
        else:
            text = extract_text_from_docx(file)
        resume_texts[os.path.basename(file)] = clean_text(text)
    print("Resume loading complete.")
    return resume_texts

# STEP 7: Optional Word Cloud (can be displayed interactively)
def generate_word_cloud(text, title="Word Cloud of Resumes"):
    if not text:
        print("No text available for word cloud generation.")
        return
    wordcloud = WordCloud(width=1000, height=500, background_color='white',
                          collocations=False, # Avoid combining common words
                          stopwords=set(WordCloud().stopwords)).generate(text)
    plt.figure(figsize=(15, 7))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title(title)
    plt.show()

# STEP 8: Main execution logic
def run_resume_matcher():
    resume_texts = load_and_clean_resumes(folder_path)
    if not resume_texts:
        print("No resumes found in the specified folder. Please upload resumes to /content/drive/MyDrive/resume data set/Resumes and try again.") # Updated message
        return

    # Generate a word cloud for all resumes
    generate_word_cloud(" ".join(resume_texts.values()), title="Overall Resume Word Cloud")

    job_description_clean = get_job_description() # Call the interactive function
    if not job_description_clean: # This check is now handled within get_job_description's retry logic
        return

    # Extract potential keywords from the job description
    # Using TF-IDF to get important words from the JD itself
    jd_vectorizer = TfidfVectorizer(stop_words='english', max_features=50) # Top 50 unique keywords
    jd_tfidf_matrix = jd_vectorizer.fit_transform([job_description_clean])
    jd_keywords = jd_vectorizer.get_feature_names_out().tolist()
    print(f"\n--- Extracted Keywords from Job Description: {', '.join(jd_keywords)} ---")

    # STEP 9: TF-IDF Matching
    # Combine JD and resumes for consistent vectorization
    all_documents = [job_description_clean] + list(resume_texts.values())
    vectorizer = TfidfVectorizer(stop_words='english') # Common English stopwords
    doc_vectors = vectorizer.fit_transform(all_documents)

    # Calculate cosine similarity of JD with all resumes
    jd_vector = doc_vectors[0:1] # The first vector is the JD
    resume_vectors = doc_vectors[1:] # The rest are resumes

    similarities = cosine_similarity(jd_vector, resume_vectors).flatten()

    # STEP 10: Ranking resumes
    ranked_resumes = sorted(zip(resume_texts.keys(), similarities), key=lambda x: x[1], reverse=True)

    print("\n" + "="*50)
    print("🔝 Top Matching Resumes:")
    print("="*50)

    # Input for number of results
    while True:
        try:
            num_results_input = input(f"How many top resumes do you want to see in the summary report? (Max {len(ranked_resumes)}): ")
            num_results = int(num_results_input)
            if 1 <= num_results <= len(ranked_resumes):
                break
            else:
                print(f"Please enter a number between 1 and {len(ranked_resumes)}.")
        except ValueError:
            print("Invalid input. Please enter a whole number.")
    num_results = min(max(1, num_results), len(ranked_resumes)) # Ensure valid range

    summary_data = []

    for i, (name, score) in enumerate(ranked_resumes[:num_results], 1):
        original_resume_text = ""
        # Need to re-read the original (uncleaned) text for highlighting context
        file_path_found = False
        for f_name, c_text in resume_texts.items():
            if f_name == name:
                full_path = os.path.join("/content/drive/MyDrive/resume data set/Resumes", f_name)
                if full_path.lower().endswith('.pdf'):
                    original_resume_text = extract_text_from_pdf(full_path)
                elif full_path.lower().endswith('.docx'):
                    original_resume_text = extract_text_from_docx(full_path)
                file_path_found = True
                break

        # If original text couldn't be loaded (e.g., due to an error), use cleaned text
        if not original_resume_text:
            original_resume_text = resume_texts[name] # Fallback to cleaned text for email/phone

        email = extract_email(original_resume_text)
        phone = extract_phone(original_resume_text)

        # Highlighted skills section using keywords extracted from JD
        highlighted_skills_html = highlight_keywords_in_context(original_resume_text, jd_keywords)

        summary_data.append({
            "Rank": i,
            "File Name": name,
            "Match Score (%)": f"{score * 100:.2f}%",
            "Email": email,
            "Phone": phone,
            "Highlighted Skills (Context)": highlighted_skills_html # Store HTML for display
        })

        print(f"\n--- Rank {i}: {name} ---")
        print(f"Match Score: {score * 100:.2f}%")
        print(f"Email: {email}")
        print(f"Phone: {phone}")

    # --- Display Summary Report using Pandas and HTML for better formatting ---
    print("\n" + "="*50)
    print("📄 Comprehensive Matching Report")
    print("="*50)

    df_summary = pd.DataFrame(summary_data)

    display_df = df_summary.drop(columns=["Highlighted Skills (Context)"]).copy()
    display(display_df.to_html(index=False), raw=True) # Display as interactive HTML table

    print("\n--- Detailed Highlighted Sections for Top Resumes ---")
    for row in summary_data:
        print(f"\n--- {row['Rank']}. {row['File Name']} (Match: {row['Match Score (%)']}) ---")
        print("\n🔍 Highlighted Skills & Context (scroll to view):")
        display(HTML(row["Highlighted Skills (Context)"]))
        print("-" * 60) # Separator for clarity

    # --- Export to CSV ---
    export_choice = input("\nDo you want to export the summary report to CSV? (yes/no): ").lower()
    if export_choice == 'yes':
        csv_filename = input("Enter filename for CSV (e.g., 'resume_matches.csv'): ")
        # Ensure 'Highlighted Skills' is saved as raw text in CSV, not HTML
        df_summary_for_csv = df_summary.copy()
        df_summary_for_csv["Highlighted Skills (Context)"] = df_summary_for_csv["Highlighted Skills (Context)"].apply(
            lambda x: re.sub(r'<[^>]*>', '', x) # Remove HTML tags for CSV
        )
        # Updated path for saving CSV
        csv_path = os.path.join("/content/drive/MyDrive/resume data set/Resumes", csv_filename)
        df_summary_for_csv.to_csv(csv_path, index=False)
        print(f"Summary exported successfully to: {csv_path}")

# Run the main function
if __name__ == '__main__':
    run_resume_matcher()

# STEP 1: Mount Google Drive
from google.colab import drive
drive.mount('/content/drive')

# STEP 2: Install packages (add fuzzywuzzy for better string matching)
!pip install python-docx PyMuPDF wordcloud scikit-learn fuzzywuzzy python-Levenshtein
!pip install --upgrade scikit-learn

# STEP 3: Imports
import os
import re
import fitz # PyMuPDF
import docx
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from fuzzywuzzy import fuzz
from IPython.display import display, HTML # For rich output in Colab

# STEP 4: Path setup (ensure this folder exists in your Google Drive)
# Create a folder named 'resume_data' inside 'My Drive' and put your resumes there
folder_path = "/content/drive/MyDrive/resume data set/Resumes" # Updated folder path
if not os.path.exists(folder_path):
    # Changed message to reflect the user's provided path
    print(f"Error: Folder not found at {folder_path}. Please ensure this folder exists in your Google Drive and contains your resumes.")
    # Exit the script if the folder doesn't exist to prevent errors
    exit()


# STEP 5: Functions
def extract_text_from_pdf(path):
    text = ""
    try:
        with fitz.open(path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF {path}: {e}")
    return text

def extract_text_from_docx(path):
    text = ""
    try:
        doc = docx.Document(path)
        text = " ".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX {path}: {e}")
    return text

def clean_text(text):
    # Remove non-ASCII characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    # Remove extra spaces and newlines
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove URLs
    text = re.sub(r'http\S+', '', text)
    # Allow only alphanumeric, common punctuation, @ and .
    text = re.sub(r'[^a-zA-Z0-9@\s.,-]', '', text)
    return text.lower()

def extract_email(text):
    match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    return match.group(0) if match else "N/A"

def extract_phone(text):
    # Broader pattern for phone numbers, including common separators and international codes
    # This pattern tries to capture 10-15 digit numbers, possibly with spaces, hyphens, or parentheses
    match = re.search(r'\+?\d{1,3}?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,9}\b', text)
    return match.group(0).replace(" ", "").replace("-", "").replace("(", "").replace(")", "") if match else "N/A"

def get_job_description():
    """Allows user to paste the job description directly with clear instructions."""
    print("\n" + "="*50)
    print("--- JOB DESCRIPTION INPUT ---")
    print("="*50)
    print("Please paste the full Job Description text below.")
    print("Type your job description line by line.")
    print("When you are finished, type 'END_JD' on a new, separate line and press Enter.")
    print("\n--- Start pasting your Job Description here: ---")

    lines = []
    while True:
        try:
            line = input() # This is where you type each line
            if line.strip().upper() == 'END_JD': # Case-insensitive check for 'END_JD'
                break
            lines.append(line)
        except EOFError:
            print("\nWarning: Received EOFError. Please ensure you type 'END_JD' on a new line to finish input.")
            # Optionally, break or return here if EOFError indicates the end of interaction in some environments
            break # Exit loop on EOFError
        except Exception as e:
            print(f"\nAn error occurred during input: {e}")
            break # Exit loop on other errors

    jd_text = "\n".join(lines)

    if not jd_text.strip(): # Check if JD is empty after stripping whitespace
        print("\nWarning: Job description was empty. Please try again.")
        # Note: Recursive call might lead to deep recursion if user repeatedly enters empty JD.
        # Consider a loop-based retry mechanism instead if this becomes an issue.
        return get_job_description() # Retry

    print("\n--- Job Description received. Processing... ---")
    return clean_text(jd_text)

def highlight_keywords_in_context(text, keywords, context_size=150):
    """Highlights keywords and provides a surrounding context."""
    highlighted_sections = []
    text_lower = text.lower() # Use lowercased text for matching

    for keyword in keywords:
        # Use a more robust pattern for whole words
        pattern = r'\b' + re.escape(keyword) + r'\b'
        for match in re.finditer(pattern, text_lower):
            start, end = match.span()
            # Find the start of the context
            context_start = max(0, start - context_size)
            # Find the end of the context
            context_end = min(len(text), end + context_size)

            # Ensure we're highlighting in the original case-preserved text
            original_match_text = text[start:end]

            # Extract the original case-preserved context
            context = text[context_start:context_end]

            # Highlight the specific matched keyword within this context for display
            # We use a non-capturing group (?:...) for the bold tag so it doesn't interfere
            highlighted_context = re.sub(r'\b' + re.escape(keyword) + r'\b', f'<b><span style="background-color: yellow;">{original_match_text}</span></b>', context, flags=re.IGNORECASE)

            highlighted_sections.append(f"...{highlighted_context}...")
    return "\n---\n".join(highlighted_sections) if highlighted_sections else "No significant skill matches found in context."


# STEP 6: Load and clean resumes
def load_and_clean_resumes(path):
    resume_files = [os.path.join(path, f) for f in os.listdir(path) if f.lower().endswith(('.pdf', '.docx'))]
    resume_texts = {}
    print(f"\n--- Loading {len(resume_files)} Resumes ---")
    for file in resume_files:
        print(f"Processing: {os.path.basename(file)}")
        if file.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file)
        else:
            text = extract_text_from_docx(file)
        resume_texts[os.path.basename(file)] = clean_text(text)
    print("Resume loading complete.")
    return resume_texts

# STEP 7: Optional Word Cloud (can be displayed interactively)
def generate_word_cloud(text, title="Word Cloud of Resumes"):
    if not text:
        print("No text available for word cloud generation.")
        return
    wordcloud = WordCloud(width=1000, height=500, background_color='white',
                          collocations=False, # Avoid combining common words
                          stopwords=set(WordCloud().stopwords)).generate(text)
    plt.figure(figsize=(15, 7))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title(title)
    plt.show()

# STEP 8: Main execution logic
def run_resume_matcher():
    resume_texts = load_and_clean_resumes(folder_path)
    if not resume_texts:
        print("No resumes found in the specified folder. Please upload resumes to /content/drive/MyDrive/resume data set/Resumes and try again.") # Updated message
        return

    # Generate a word cloud for all resumes
    generate_word_cloud(" ".join(resume_texts.values()), title="Overall Resume Word Cloud")

    job_description_clean = get_job_description() # Call the interactive function
    if not job_description_clean: # This check is now handled within get_job_description's retry logic
        return

    # Extract potential keywords from the job description
    # Using TF-IDF to get important words from the JD itself
    jd_vectorizer = TfidfVectorizer(stop_words='english', max_features=50) # Top 50 unique keywords
    jd_tfidf_matrix = jd_vectorizer.fit_transform([job_description_clean])
    jd_keywords = jd_vectorizer.get_feature_names_out().tolist()
    print(f"\n--- Extracted Keywords from Job Description: {', '.join(jd_keywords)} ---")

    # STEP 9: TF-IDF Matching
    # Combine JD and resumes for consistent vectorization
    all_documents = [job_description_clean] + list(resume_texts.values())
    vectorizer = TfidfVectorizer(stop_words='english') # Common English stopwords
    doc_vectors = vectorizer.fit_transform(all_documents)

    # Calculate cosine similarity of JD with all resumes
    jd_vector = doc_vectors[0:1] # The first vector is the JD
    resume_vectors = doc_vectors[1:] # The rest are resumes

    similarities = cosine_similarity(jd_vector, resume_vectors).flatten()

    # STEP 10: Ranking resumes
    ranked_resumes = sorted(zip(resume_texts.keys(), similarities), key=lambda x: x[1], reverse=True)

    print("\n" + "="*50)
    print("🔝 Top Matching Resumes:")
    print("="*50)

    # Input for number of results
    while True:
        try:
            num_results_input = input(f"How many top resumes do you want to see in the summary report? (Max {len(ranked_resumes)}): ")
            num_results = int(num_results_input)
            if 1 <= num_results <= len(ranked_resumes):
                break
            else:
                print(f"Please enter a number between 1 and {len(ranked_resumes)}.")
        except ValueError:
            print("Invalid input. Please enter a whole number.")
    num_results = min(max(1, num_results), len(ranked_resumes)) # Ensure valid range


    summary_data = []

    for i, (name, score) in enumerate(ranked_resumes[:num_results], 1):
        original_resume_text = ""
        # Need to re-read the original (uncleaned) text for highlighting context
        file_path_found = False
        for f_name, c_text in resume_texts.items():
            if f_name == name:
                full_path = os.path.join("/content/drive/MyDrive/resume data set/Resumes", f_name)
                if full_path.lower().endswith('.pdf'):
                    original_resume_text = extract_text_from_pdf(full_path)
                elif full_path.lower().endswith('.docx'):
                    original_resume_text = extract_text_from_docx(full_path)
                file_path_found = True
                break

        # If original text couldn't be loaded (e.g., due to an error), use cleaned text
        if not original_resume_text:
            original_resume_text = resume_texts[name] # Fallback to cleaned text for email/phone


        email = extract_email(original_resume_text)
        phone = extract_phone(original_resume_text)

        # Highlighted skills section using keywords extracted from JD
        highlighted_skills_html = highlight_keywords_in_context(original_resume_text, jd_keywords)

        summary_data.append({
            "Rank": i,
            "File Name": name,
            "Match Score (%)": f"{score * 100:.2f}%",
            "Email": email,
            "Phone": phone,
            "Highlighted Skills (Context)": highlighted_skills_html # Store HTML for display
        })

        print(f"\n--- Rank {i}: {name} ---")
        print(f"Match Score: {score * 100:.2f}%")
        print(f"Email: {email}")
        print(f"Phone: {phone}")


    # --- Display Summary Report using Pandas and HTML for better formatting ---
    print("\n" + "="*50)
    print("📄 Comprehensive Matching Report")
    print("="*50)

    df_summary = pd.DataFrame(summary_data)

    # For displaying in Colab, we can render the DataFrame as HTML
    # We'll need to handle the 'Highlighted Skills' column carefully
    # Creating a temporary DataFrame for display that shows raw text,
    # and then separately displaying the detailed HTML for each top resume.

    display_df = df_summary.drop(columns=["Highlighted Skills (Context)"]).copy()
    display(display_df.to_html(index=False), raw=True) # Display as interactive HTML table

    print("\n--- Detailed Highlighted Sections for Top Resumes ---")
    for row in summary_data:
        print(f"\n--- {row['Rank']}. {row['File Name']} (Match: {row['Match Score (%)']}) ---")
        print("\n🔍 Highlighted Skills & Context (scroll to view):")
        display(HTML(row["Highlighted Skills (Context)"]))
        print("-" * 60) # Separator for clarity


    # --- Export to CSV ---
    export_choice = input("\nDo you want to export the summary report to CSV? (yes/no): ").lower()
    if export_choice == 'yes':
        csv_filename = input("Enter filename for CSV (e.g., 'resume_matches.csv'): ")
        # Ensure 'Highlighted Skills' is saved as raw text in CSV, not HTML
        df_summary_for_csv = df_summary.copy()
        df_summary_for_csv["Highlighted Skills (Context)"] = df_summary_for_csv["Highlighted Skills (Context)"].apply(
            lambda x: re.sub(r'<[^>]*>', '', x) # Remove HTML tags for CSV
        )
        # Updated path for saving CSV
        csv_path = os.path.join("/content/drive/MyDrive/resume data set/Resumes", csv_filename)
        df_summary_for_csv.to_csv(csv_path, index=False)
        print(f"Summary exported successfully to: {csv_path}")

# Run the main function
if __name__ == '__main__':
    run_resume_matcher()